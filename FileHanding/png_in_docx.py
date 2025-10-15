import os
import random
from pathlib import Path

from docx import Document
from docx.shared import Inches

docx_dir = "/Users/lyu/Code/GitHub/PythonScripts/docx_out"
png_dir = "/Users/lyu/Code/GitHub/PythonScripts/out"
output_dir = "/Users/lyu/Code/GitHub/PythonScripts/new_out"
log_file = "/Users/lyu/Code/GitHub/PythonScripts/new_out/日志.txt"
insert_width_inches = 2.0

docx_dir_path = Path(docx_dir)
png_dir_path = Path(png_dir)
output_dir_path = Path(output_dir)
log_file_path = Path(log_file)

output_dir_path.mkdir(parents=True, exist_ok=True)

# 获取所有 DOCX 和 PNG 文件，并按名字排序（保证顺序一致）
docx_files = sorted(docx_dir_path.glob("*.docx"))
png_files = sorted(png_dir_path.glob("*.png"))

if not docx_files:
    raise ValueError("未找到 docx 文件")
if not png_files:
    raise ValueError("未找到 PNG 文件")

rng = random.Random()
shuffled_docx = docx_files[:]
rng.shuffle(shuffled_docx)
shuffled_png = png_files[:]
rng.shuffle(shuffled_png)

log_records = []

# 遍历 DOCX 文件，随机插入 PNG（每张图片仅使用一次）
for docx_path, png_path in zip(shuffled_docx, shuffled_png):
    doc = Document(docx_path)

    position_desc = None
    paragraphs = doc.paragraphs
    width = Inches(insert_width_inches)

    if paragraphs:
        idx = rng.randrange(len(paragraphs))
        paragraph = paragraphs[idx]
        if rng.choice([True, False]):
            target_paragraph = paragraph.insert_paragraph_before()
            run = target_paragraph.add_run()
            position_desc = f"before paragraph {idx}"
        else:
            run = paragraph.add_run()
            position_desc = f"in paragraph {idx}"
        run.add_picture(str(png_path), width=width)
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(str(png_path), width=width)
        position_desc = "in new paragraph at document end"

    # 保存新文件
    output_path = output_dir_path / docx_path.name
    doc.save(output_path)

    log_entry = f"{docx_path.name} <- {png_path.name} @ {position_desc}"
    log_records.append(log_entry)
    print(f"已处理 {log_entry}")

# 保存日志
log_file_path.write_text("\n".join(log_records), encoding="utf-8")

print(f"\n所有 docx 已随机插入签名并保存到 {output_dir_path}")
print(f"插入日志已保存到 {log_file_path}")
